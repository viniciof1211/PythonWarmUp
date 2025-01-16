from docx import Document
from docx.shared import Pt

# Create a new Word document
doc = Document()

# Add Title
doc.add_heading("Ta_amei Mikra: Harmony in Chanting, Coding, and Real-Life Algorithms", level=1)

# Set Times New Roman font for the whole document
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(22)

# Data for Ta'amei Mikra
data = {
    "Ta_am Name": [
        "Etnachta ֑", "Sof Pasuk ֽ", "Zakef Gadol ֔", "Zakef Katan ֕", "Segol ֡", "Revia ֗",
        "Munach ֣", "Tipcha ֛", "Mercha ֥", "Kadma ֨", "Azla ֦", "Shalshelet ֓", "Pashta ֙",
        "Tlisha Gedola ֜", "Tlisha Ketana ֩", "Yerach Ben Yomo ֬", "Pazer ֙", "Telisha Ketana ֭"
    ],
    "Chanting Function": [
        "A main pause, halfway through a verse, like a semicolon.",
        "Marks the end of a verse, similar to a period in punctuation.",
        "A strong pause or emphasis within a clause.",
        "A shorter pause with mild emphasis.",
        "Groups words together, creating a melodic pattern.",
        "Signals a slight pause or a mid-clause pause.",
        "Smooth transition note; connects phrases smoothly.",
        "Indicates a preparatory pause before a more significant note.",
        "A smooth connecting note leading to the next phrase.",
        "Forward motion or progression within a sentence.",
        "Forward progression, often with momentum.",
        "A rare, dramatic pause with significant emphasis.",
        "Used to mark a mild separation, often before repetition.",
        "An ornamental flourish on a word.",
        "Short decorative pause, very light emphasis.",
        "Indicates a smooth or connecting tone with a rare appearance.",
        "Indicates a complex or long grouping of words.",
        "Very light pause or ornamental flourish on minor phrases."
    ],
    "Coding Metaphor": [
        "Base case of recursion: Stops the recursion midway for processing.",
        "End of recursion or algorithm execution (termination state).",
        "Pivot selection in Divide and Conquer (e.g., Quicksort).",
        "Checkpoint in iterative processes or loops.",
        "Logical grouping, similar to partitioning data (e.g., Mergesort).",
        "Midway processing or marking dependencies.",
        "Smooth transitions in algorithm steps, like chaining operations.",
        "Pre-processing phase before a significant algorithmic operation.",
        "Intermediate step in processes (e.g., BFS/DFS transitions).",
        "Indicates forward recursion or step progression.",
        "Momentum-based decisions, like Binary Search mid-point.",
        "Rare critical state, like a deadlock detection or fallback condition.",
        "Light separation of repeated steps in loops.",
        "Optimization or embellishment for minor improvements.",
        "Micro-checkpoints, like inline optimizations.",
        "Uncommon transitions between linked nodes or edges.",
        "Complex clustering, like graph partitioning.",
        "Lightweight optional check in conditional logic."
    ],
    "Real-Life Example": [
        "Base case in Mergesort or Quicksort: halting further recursion.",
        "Termination of a program or clean return in recursion.",
        "Quicksort pivot selection determining split logic.",
        "Checkpoint in for-loops or while-loops, like an iterator break.",
        "Grouping nodes in graph algorithms for SCC (Strongly Connected Components).",
        "Mid-point handling in Binary Search or dynamic programming.",
        "Linked list traversal or smooth transitions in Hash Map chaining.",
        "DFS pre-processing before visiting child nodes.",
        "In-order traversal step in Binary Tree algorithms.",
        "Forward recursive call in Quicksort.",
        "Binary Search: Moving forward to the next half.",
        "Rare state detection in Deadlock or cycle detection in DFS.",
        "Repeated elements in loops for cache optimizations.",
        "Adding lightweight optimizations for decorators in Python.",
        "Inline if-statements or micro-adjustments for efficiency.",
        "Edge case detection between two states in BFS or graphs.",
        "Graph clustering using community detection algorithms.",
        "Optional step checks in decision trees or fuzzy logic."
    ]
}

def generate_cantillations():
    # Add table headers
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Ta_am Name"
    hdr_cells[1].text = "Chanting Function"
    hdr_cells[2].text = "Coding Metaphor"
    hdr_cells[3].text = "Real-Life Example"

    # Populate the table
    for index in range(len(data["Ta_am Name"])):
        row_cells = table.add_row().cells
        row_cells[0].text = data['Ta_am Name'][index]
        row_cells[1].text = data['Chanting Function'][index]
        row_cells[2].text = data['Coding Metaphor'][index]
        row_cells[3].text = data['Real-Life Example'][index]

    # Save the Word document
    file_name = "Taamei_Mikra_Algorithms.docx"
    doc.save(file_name)
    print(f"Document saved successfully as {file_name}")
